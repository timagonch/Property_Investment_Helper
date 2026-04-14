"""
Generate a comprehensive project reference document (Word .docx).
Run with: uv run python scripts/build_project_doc.py
Output:   documentation/Project_Reference_Guide.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, os

BASE = Path(__file__).resolve().parent.parent
OUT  = BASE / "documentation" / "Project_Reference_Guide.docx"
OUT.parent.mkdir(exist_ok=True)

# ── Colour palette ────────────────────────────────────────────────────────────
RED    = RGBColor(0xE6, 0x39, 0x46)
NAVY   = RGBColor(0x0D, 0x1B, 0x2A)
GRAY   = RGBColor(0x6C, 0x75, 0x7D)
BLACK  = RGBColor(0x21, 0x25, 0x29)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Helper functions ──────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(18)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RED
    p.runs[0].font.size = Pt(13)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(11)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10.5)
    p.runs[0].font.color.rgb = BLACK
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.runs[0].font.size = Pt(10.5)
    return p

def decision_box(title, reasoning, implication):
    """Highlighted decision block."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    run = p.add_run(f"  DECISION: {title}  ")
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)
    # shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "0D1B2A")
    pPr.append(shd)

    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent  = Inches(0.3)
    p2.paragraph_format.right_indent = Inches(0.3)
    r1 = p2.add_run("Why: ")
    r1.bold = True; r1.font.size = Pt(10)
    r2 = p2.add_run(reasoning)
    r2.font.size = Pt(10)

    p3 = doc.add_paragraph()
    p3.paragraph_format.left_indent  = Inches(0.3)
    p3.paragraph_format.right_indent = Inches(0.3)
    r3 = p3.add_run("Implication: ")
    r3.bold = True; r3.font.size = Pt(10); r3.font.color.rgb = RED
    r4 = p3.add_run(implication)
    r4.font.size = Pt(10)
    doc.add_paragraph()

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0D1B2A")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        fill  = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9.5)
            tc = cells[c_idx]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run("NC/SC Neighborhood Market\nTransition Monitor")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = NAVY

doc.add_paragraph()
t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t2.add_run("Project Reference Guide — Spring 2026")
r.font.size = Pt(12)
r.font.color.rgb = GRAY

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t3.add_run("Business Intelligence Group Project")
r2.font.size = Pt(11)
r2.font.color.rgb = GRAY

doc.add_paragraph()
t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t4.add_run("Live App:  https://huggingface.co/spaces/timagonch/nc-sc-market-monitor")
r3.font.size = Pt(10)
r3.font.color.rgb = RED
r3.font.bold = True

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Executive Summary")
body(
    "This project is an early-warning system for neighborhood housing market transitions across "
    "North Carolina and South Carolina. It identifies ZIP codes that are showing leading signals "
    "of sustained price momentum shifts — upward or downward — months before those shifts appear "
    "in widely-reported headline metrics."
)
body(
    "The core output is a per-ZIP opportunity score: the model-predicted probability that a ZIP "
    "will enter an upward housing transition over the next 12 months. A companion directional "
    "score (P(up) − P(down)) separates heating from cooling markets. Both are available at the "
    "monthly level for 655 NC/SC ZIP codes, with validated predictions for 2023 and an "
    "validated 2025 forward signal."
)
body(
    "The model is a Gradient Boosting classifier trained and evaluated using strict walk-forward "
    "validation (no random splits on time-series data). The winning model uses only Zillow and "
    "Redfin monthly data — 18 features — and achieves AUC 0.865 on the 2023 held-out test set. "
    "The 2024 predictions were subsequently validated against actual 2025 outcomes (AUC 0.7747), "
    "and the model was retrained on Jun 2019–Dec 2024 to generate the current 2025 forward signal."
)

h2("Key Numbers at a Glance")
table(
    ["Metric", "Value"],
    [
        ["ZIP codes covered",          "655 (NC and SC)"],
        ["Training data window",       "Jun 2019 – Dec 2024 (retrained model)"],
        ["Model",                      "Gradient Boosting — Monthly-Only (Zillow + Redfin)"],
        ["Features selected (RFECV)",  "18"],
        ["Validation approach",        "Walk-forward — 5 folds / evaluation points"],
        ["Primary AUC (test 2023)",    "0.865"],
        ["2024 predictions vs 2025 actuals", "AUC 0.7747 (retrospective validation)"],
        ["Forward signal available",   "Dec 2025 – Nov 2026 (from Nov 2025 features)"],
        ["Live app",                   "huggingface.co/spaces/timagonch/nc-sc-market-monitor"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM FRAMING
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. Problem Framing")
body(
    "A neighborhood transition is a sustained, directional shift in a ZIP code's housing market "
    "dynamics — typically visible in momentum data months before it shows up in prices. By the "
    "time median sale price in a neighborhood has clearly moved, the opportunity window for "
    "investors has often already closed."
)
body(
    "The goal is to detect early-stage transition signals using leading indicators (inventory "
    "trends, days-on-market acceleration, above-list sale rates) rather than lagging price "
    "metrics alone."
)

h2("The Four Core Questions")
bullet("Which ZIPs are showing the strongest signals of emerging transition pressure?")
bullet("Are there distinct neighborhood market archetypes (structural market types) across NC/SC?")
bullet("Will a ZIP enter an upward or downward transition within the next 12 months?")
bullet("Which leading indicators drive that prediction most, and how much does each contribute?")

h2("Terminology — Important Distinctions")
body(
    "These terms are used precisely throughout the project and the app. Using them correctly "
    "matters especially in a finance/investment context."
)
table(
    ["Term", "Definition", "Why it matters"],
    [
        ["Opportunity score",
         "Predicted probability of upward transition (binary model output). Range 0–1.",
         "One-sided — only measures upward probability. A low score doesn't tell you if a market is cooling vs. stable."],
        ["Net direction score",
         "P(up) − P(down) from the 3-class model. Range −1 to +1.",
         "Two-sided — distinguishes heating, stable, and cooling. Use this to spot ZIPs to avoid, not just ZIPs to buy."],
        ["Cooling risk",
         "Downward transition probability — the blue end of the direction map.",
         "We never call upward transition probability 'risk'. In finance, risk implies downside. High opportunity = buy signal, not a warning."],
        ["Transition",
         "A sustained, directional shift in market dynamics lasting 12+ months.",
         "Not a one-month spike. The target is engineered over a forward 12-month window, not a single data point."],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. DATA SOURCES
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Data Sources")
body(
    "Six sources were integrated. The final winning model uses only the two monthly sources "
    "(Zillow + Redfin). The annual sources were evaluated but dropped — see Section 6 for the "
    "reasoning."
)
table(
    ["Source", "Frequency", "What it provides", "In final model?"],
    [
        ["Zillow Research (ZHVI)", "Monthly", "Home Value Index by ZIP code", "Yes"],
        ["Redfin Market Tracker",  "Monthly", "Sale/list price, inventory, DOM, sale-to-list ratio, homes sold, off-market speed", "Yes"],
        ["ACS 5-year estimates",   "Annual",  "Median income, rent, home value, owner occupancy rate", "No — dropped"],
        ["County Business Patterns (CBP)", "Annual", "Establishments, employment, annual payroll", "No — dropped"],
        ["IRS Statistics of Income (SOI)", "Annual", "ZIP-level income and tax return indicators", "No — dropped"],
        ["Freddie Mac PMMS",       "Weekly",  "30yr/15yr/ARM mortgage rates", "No — excluded early"],
    ]
)

h2("Why Mortgage Rates Were Excluded Early")
decision_box(
    "Remove Freddie Mac mortgage rate features before modeling",
    "Mortgage rates are national — they have zero variance across ZIP codes within any given "
    "month. Including them lets the model learn that 'when rates are low, transitions happen', "
    "which is a temporal trend, not a ZIP-specific signal. In V1 of the model, "
    "mortgage_rate_30yr was the single most important feature — which was a red flag, not a "
    "success. A model that predicts based on national rate levels cannot help identify which "
    "specific ZIPs to target.",
    "All mortgage features were removed before any model comparison. This dropped V1 AUC from "
    "0.861 to 0.847 — a small and acceptable cost for a model that actually captures ZIP-level "
    "differentiation."
)

h2("Why Annual Sources Were Dropped from the Final Model")
decision_box(
    "Drop ACS, CBP, and IRS annual features from the winning model",
    "After RFECV feature selection on all sources combined (V3 model), the 17 annual features "
    "ranked #12 through #27 in permutation importance (0.0008–0.0024 vs. 0.046 for the top "
    "feature). They were providing almost no marginal information beyond what the monthly "
    "features already captured. More importantly, annual data has a structural lag — ACS "
    "estimates for year Y are published in year Y+1, meaning the model would be using "
    "information that wasn't actually available at prediction time in real deployment.",
    "The monthly-only model (V3 without annual features) achieved higher AUC (0.865 vs 0.852) "
    "than V3 with all sources, and higher F1. Removing the annual features forced RFECV to "
    "select better monthly substitutes and eliminated the data-lag problem. It also unlocks "
    "future extension: Redfin data is available back to 2012, meaning the model could be "
    "retrained with 4x more data and validated on 2023/2024 without any annual survey dependency."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. ZIP COVERAGE DECISIONS
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. ZIP Code Coverage Decisions")
body(
    "The raw dataset began with all NC/SC ZIP codes in the Zillow universe. Several filtering "
    "steps reduced and then partially recovered this set. Understanding these steps is important "
    "because they define which ZIPs the model can speak to — and which it cannot."
)

h2("Step 1 — Redfin Coverage Filter")
body(
    "Redfin's Market Tracker only covers ZIPs with sufficient transaction volume to report "
    "reliable monthly statistics. ZIPs with no Redfin coverage at all were removed from the "
    "universe entirely."
)
decision_box(
    "Remove ZIPs with no Redfin coverage",
    "The model relies heavily on Redfin features (inventory trends, DOM, sale-to-list ratios). "
    "A ZIP with no Redfin data cannot be scored by the model.",
    "Reduced the universe from all NC/SC ZIPs to 675 ZIPs with at least some Redfin coverage."
)

h2("Step 2 — Continuity Filter (Initial: Too Strict)")
body(
    "After Redfin filtering, ZIPs with long gaps (missing streaks) in any of the model feature "
    "columns were removed. The initial implementation checked all 39 model columns — meaning "
    "a ZIP was dropped if it had gaps in any feature, including annual ACS/IRS features which "
    "legitimately have gaps by design."
)
decision_box(
    "Relax the streak filter from all 39 columns to the 5 base columns only",
    "The original 39-column check was dropping ZIPs that had complete monthly data but "
    "incidental gaps in annual survey columns (ACS, CBP, IRS). These annual gaps are not "
    "real data quality problems — they're an artifact of how annual data is joined to monthly "
    "panels. Since the winning model doesn't use annual features anyway, these ZIPs should "
    "not have been excluded.",
    "Recovered 115 additional ZIPs — expanding coverage from 540 to 655 ZIPs. The 5 base "
    "columns checked are: zhvi (Zillow), HOMES_SOLD, INVENTORY_mom_pct, INVENTORY_yoy_pct, "
    "and B25077_001E (ACS median owner home value — one annual feature kept as a baseline "
    "context variable). 83 ZIPs with genuine monthly data gaps remain excluded."
)

h2("Step 3 — GeoJSON Map Update")
body(
    "The choropleth maps in the app use a GeoJSON file of Census ZCTA (ZIP Code Tabulation "
    "Area) boundaries. The original GeoJSON was built for 540 ZIPs. After expanding to 655, "
    "the map was silently missing 119 ZIPs — they had model predictions but no geometry, so "
    "they were invisible on the map."
)
decision_box(
    "Rebuild nc_sc_zips.geojson from the Census 2020 500k cartographic boundary file",
    "The existing GeoJSON was pre-filtered to the original 540 ZIPs and couldn't be updated "
    "by hand. The Census 500k shapefile was already downloaded locally, so we used pyshp to "
    "extract and convert the boundaries for all 655 model ZIPs to GeoJSON.",
    "The GeoJSON grew from 540 features (629KB) to 655 features (4MB). All ZIPs are now "
    "visible on the map. The 4MB size is still manageable for Plotly choropleth rendering."
)

h2("Final ZIP Count Summary")
table(
    ["Stage", "ZIPs", "Reason for change"],
    [
        ["All NC/SC ZIPs in Zillow universe",     "~800+", "Starting point"],
        ["After Redfin coverage filter",           "675",   "Removed ZIPs with no Redfin data at all"],
        ["After original 39-column streak filter", "540",   "Removed ZIPs with gaps in any model column (too strict)"],
        ["After relaxed 5-column streak filter",   "655",   "Recovered ZIPs with gaps only in annual columns"],
        ["ZIPs still excluded",                    "20",    "83 ZIPs with genuine monthly data gaps remain out"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. TARGET VARIABLE CONSTRUCTION
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Target Variable Construction")
body(
    "One of the most consequential design decisions in the project was how to define what we "
    "are trying to predict. There is no pre-labeled dataset of 'neighborhood transitions' — we "
    "had to engineer the target from raw market data."
)

h2("The Composite Transition Score")
body(
    "We defined a neighborhood transition as a sustained, directional shift in three "
    "complementary market signals measured over the forward 12-month window:"
)
table(
    ["Signal", "How measured", "Direction of transition signal"],
    [
        ["Home value growth",      "Forward 12-month avg YoY growth minus trailing 12-month baseline (per ZIP)", "Higher than baseline = positive"],
        ["Inventory tightening",   "Forward 12-month avg MoM inventory change, negated",                        "Falling supply = positive (negated so tightening = +)"],
        ["Market competitiveness", "Forward 12-month avg % of homes sold above list price",                     "More above-list sales = positive"],
    ]
)
body(
    "Each signal is z-scored across all ZIPs within the same month (mean 0, std 1), then "
    "the three z-scores are equally weighted and averaged to produce the composite score. "
    "Z-scoring ensures that no single signal dominates due to scale differences, and that "
    "each ZIP is evaluated relative to the cross-sectional distribution of that month — "
    "not against an absolute threshold."
)

decision_box(
    "Use z-scored composite score rather than a single price metric",
    "Relying on price growth alone misses the early signals. Inventory tightening and "
    "above-list sales typically precede price appreciation by several months. By combining "
    "all three, the composite captures the full transition dynamic — and a ZIP that scores "
    "high on all three is more convincingly transitioning than one with a price spike alone.",
    "The composite score is a forward-looking measure computed from the 12 months AFTER "
    "the observation date. This means the model is trained to predict what this score will "
    "look like in the future, using only past and current features as inputs."
)

h2("From Score to Labels")
body(
    "The composite score is continuous. We convert it to classification targets using "
    "percentile thresholds computed on the final 16,740-row modeling subset:"
)
table(
    ["Target", "Type", "Label rule", "Class balance"],
    [
        ["transition_next_12m",  "Binary",  "1 if composite score ≥ 75th percentile (upward transition)",          "25% positive / 75% negative"],
        ["transition_direction", "3-class", "+1 if ≥ p75 (up) / 0 if p25–p75 (stable) / −1 if < p25 (down)",    "25% up / 50% stable / 25% down"],
    ]
)
decision_box(
    "Use percentile thresholds rather than absolute values",
    "Percentile thresholds automatically adjust for market-wide trends. If the entire "
    "housing market surged (as in 2020–2021), using an absolute threshold would label "
    "nearly everything as 'transitioning'. Percentiles ensure the labels always reflect "
    "relative outperformers — which is what an investor actually cares about.",
    "The 25/50/25 class split gives a clean and interpretable distribution. The binary "
    "target (transition_next_12m) is the primary prediction target. The 3-class target "
    "adds directionality — it can identify ZIPs likely to decline, which the binary model "
    "cannot distinguish from stable ZIPs."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. FEATURE ENGINEERING
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Feature Engineering")
body(
    "Raw Zillow and Redfin data provides level measurements each month (e.g. home value index, "
    "inventory count). To give the model signal about momentum, acceleration, and divergence "
    "from historical norms, we engineered several families of derived features."
)

h2("Feature Families")
table(
    ["Family", "Examples", "What it captures"],
    [
        ["Month-over-month change",   "home_value_mom_pct, inventory_mom_pct",              "Current momentum — is the market moving and in which direction?"],
        ["Rolling averages",          "home_value_mom_3m_avg, home_value_mom_12m_avg, inventory_mom_12m_avg", "Smoothed trend — removes noise from a single month's reading"],
        ["Acceleration",              "home_value_accel",                                   "Rate of change of the rate of change — is momentum itself speeding up or slowing down?"],
        ["Year-over-year change",     "home_value_yoy_pct, inventory_yoy_pct",              "Medium-term direction vs. same month last year"],
        ["Baseline deviation",        "home_value_vs_baseline, dom_vs_baseline",            "Is this ZIP behaving unusually vs. its own historical norm?"],
        ["Lagged features",           "home_value_mom_pct_lag6, pct_sold_above_list_lag6, home_value_vs_baseline_lag6, avg_sale_to_list_ratio_lag6", "Recent history — what was the market doing 6 months ago?"],
        ["Level features",            "home_value_index, homes_sold, inventory, new_listings, median_days_on_market, pct_off_market_in_2wks", "Absolute market conditions at this point in time"],
    ]
)

h2("Why Lags Matter")
body(
    "The 6-month lag features deserve special explanation. The model is predicting what will "
    "happen over the NEXT 12 months. Features computed from current data are the most recent "
    "signal, but 6-month lagged features capture whether the market was already building "
    "momentum before today. A ZIP where pct_sold_above_list was high 6 months ago AND is "
    "still high today has sustained demand pressure — a stronger signal than a recent spike."
)
body(
    "SHAP analysis confirms this intuition: home_value_vs_baseline_lag6 is the second most "
    "important feature in the model (mean |SHAP| = 0.73), behind only home_value_mom_pct "
    "(mean |SHAP| = 1.14)."
)

h2("The 18 Final RFECV-Selected Features")
table(
    ["Feature", "Category", "SHAP rank (approx)"],
    [
        ["home_value_mom_pct",          "MoM momentum",   "#1"],
        ["home_value_vs_baseline_lag6", "Lagged baseline", "#2"],
        ["home_value_accel",            "Acceleration",    "#3"],
        ["inventory_mom_12m_avg",       "Supply trend",    "#4"],
        ["pct_sold_above_list_lag6",    "Lagged demand",   "#5"],
        ["avg_sale_to_list_ratio_lag6", "Lagged demand",   "#6"],
        ["home_value_vs_baseline",      "Baseline dev.",   "#7"],
        ["home_value_mom_3m_avg",       "Smoothed MoM",    "#8"],
        ["home_value_mom_12m_avg",      "Long-run MoM",    "#9"],
        ["home_value_mom_pct_lag6",     "Lagged MoM",      "#10"],
        ["home_value_yoy_pct",          "YoY growth",      "#11"],
        ["inventory_yoy_pct",           "Supply YoY",      "#12"],
        ["dom_vs_baseline",             "DOM deviation",   "#13"],
        ["homes_sold",                  "Activity level",  "#14"],
        ["new_listings",                "Supply flow",     "#15"],
        ["inventory",                   "Supply level",    "#16"],
        ["median_days_on_market",       "Demand speed",    "#17"],
        ["pct_off_market_in_2wks",      "Demand urgency",  "#18"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. CLUSTERING — NEIGHBORHOOD ARCHETYPES
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Clustering — Neighborhood Archetypes")
body(
    "Before building the predictive model, we used unsupervised learning to characterize the "
    "structural market identity of each ZIP code. This is separate from the prediction — "
    "clustering answers 'what kind of market is this ZIP fundamentally?' while the model "
    "answers 'where is it heading next?'"
)

h2("Method")
body(
    "K-means clustering on the same 18 features selected by RFECV for the prediction model "
    "(means computed over Jun 2019–Dec 2024). Using the model's own feature set improved "
    "silhouette from 0.11 to 0.16 vs. the original 22-feature approach. Each ZIP gets one "
    "cluster assignment representing its long-run market character."
)
decision_box(
    "Use static (time-invariant) cluster assignments rather than monthly re-clustering",
    "The clusters are meant to characterize a ZIP's structural identity — its long-run "
    "market type — not its current momentum. The prediction model already captures the "
    "time-varying dynamics. Having both time-varying clusters and time-varying predictions "
    "would be redundant and harder to interpret. A ZIP being 'Affordable High Appreciation' is a "
    "durable property of its market structure; whether it's about to break out is what "
    "the model predicts.",
    "Cluster labels are stable context for interpretation, not additional prediction signals."
)

h2("Silhouette Score and Honest Interpretation")
decision_box(
    "Report and acknowledge the silhouette score of 0.16",
    "K-means silhouette scores in geographic/economic clustering are commonly low because "
    "real markets exist on a continuum rather than in discrete clusters. A score of 0.16 "
    "does not mean the clustering is wrong — it means the boundaries are fuzzy, which is "
    "an accurate reflection of reality. Using the same 18 features as the prediction model "
    "improved separation (0.16 vs. 0.11) and ensures the archetypes are defined in the same "
    "language the model speaks.",
    "The clusters should be presented as 'market tendencies on a spectrum', not as hard, "
    "mutually exclusive categories. The profiles are real and interpretable — a ZIP in the "
    "'Affordable High Appreciation' cluster genuinely behaves differently than one in "
    "'High-Value, Tight Supply' — but the boundary between clusters is not sharp."
)

h2("The 5 Archetypes")
table(
    ["Archetype", "ZIPs", "Key characteristics", "Avg opportunity score"],
    [
        ["💎 High-Value, Tight Supply",       "74",  "Highest median values (~$318K), only cluster with declining inventory (−8% YoY), high off-market rate (41%). Durable demand absorbs supply consistently.", "~45%"],
        ["🧊 Low-Activity, Cooling",          "44",  "Lowest values (~$162K), weakest appreciation (7.1% YoY), highest inventory growth (+14% YoY), slowest sales pace (76-day DOM). Softest market on all metrics.", "~25%"],
        ["🔥 Competitive Mid-Market",         "211", "Mid-range values (~$266K), highest above-list rate (37.6%), fastest DOM (45 days). Strongest competition of any cluster — broad geographic representation.", "~52%"],
        ["⚡ Affordable, High Appreciation",  "172", "Entry-level (~$216K) with strongest YoY appreciation (12.8%). Rising inventory signals growing developer and investor interest. Best affordability-momentum combination.", "~55%"],
        ["📦 Moderate Value, Supply Growing", "154", "Mid-range (~$213K), rising inventory (+8% YoY), slow DOM (76 days), softest competition. Supply outpacing demand — appreciation is holding but buyers have options.", "~38%"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. SUPERVISED MODELING
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Supervised Modeling")

h2("Why Walk-Forward Validation — Not Random Splits")
decision_box(
    "Use walk-forward (time-ordered) validation exclusively",
    "The data is a time-series panel: the same 655 ZIPs observed monthly over several years. "
    "Random train/test splits would allow the model to train on June 2021 data and test on "
    "January 2021 data — learning from the future to predict the past. This leaks information "
    "and produces wildly optimistic AUC scores that won't hold up in real deployment. Walk-forward "
    "training matches how the model would actually be used: always trained on the past, "
    "predicting the future.",
    "All reported AUC and F1 scores are from walk-forward evaluation only. Any model evaluated "
    "with random splits on this dataset should be considered unreliable."
)

h2("Walk-Forward Folds and Retrospective Validation")
table(
    ["Fold", "Training data", "Test / Evaluation", "AUC", "Note"],
    [
        ["A",            "Jun 2019 – Dec 2020", "Jan–Dec 2021",        "0.8443", "COVID-era hot market — strong signal"],
        ["B",            "Jun 2019 – Dec 2021", "Jan–Dec 2022",        "0.8153", "Fed rate-hike year — market inflection, hardest fold"],
        ["C (primary)",  "Jun 2019 – Dec 2022", "Jan–Dec 2023",        "0.8652", "Post-normalization — model at its best with most training data"],
        ["D (retrain)",  "Jun 2019 – Dec 2023", "Jan–Dec 2024",        "0.7240", "Retrained model — powers the 2025 forward signal"],
        ["Actuals",      "Jun 2019 – Dec 2022", "2025 real outcomes",  "0.7747", "2024 predictions vs. actual 2025 transitions — retrospective validation"],
    ]
)
body(
    "The 2022 AUC dip to 0.815 is expected and not a failure. The Fed began an aggressive "
    "rate-hike cycle in early 2022, causing the sharpest housing market correction in decades. "
    "The model was trained on 2019–2021 (pre-hike conditions) and tested on a completely "
    "different regime. The 2023 recovery to 0.865 shows the model captures durable structural "
    "signals, not just COVID-era noise. The retrospective 2025 validation (AUC 0.7747) further "
    "confirms the model generalises across multiple market cycles."
)

h2("Model Selection Process")
table(
    ["Version", "Description", "AUC (Fold B)", "Decision"],
    [
        ["V1", "All features including mortgage rates",        "0.861", "Mortgage rate was #1 feature — national signal, not ZIP-specific. Rejected."],
        ["V2", "V1 minus mortgage features",                   "0.847", "ΔAUC = −0.013 (acceptable). Baseline for fair comparison."],
        ["V3", "RFECV on all sources (monthly + annual)",      "0.852", "39 features selected. Annual features ranked #12–#27. Kept for comparison."],
        ["Monthly-Only ✓", "RFECV on Zillow + Redfin only",  "0.865", "18 features. Beats V3 on both AUC and F1. Selected as winner."],
    ]
)

h2("RFECV Feature Selection")
body(
    "Recursive Feature Elimination with Cross-Validation (RFECV) was used to identify the "
    "minimum feature set that preserves model performance. RFECV fits the model repeatedly, "
    "removing the least important feature each round, and uses cross-validation to measure "
    "whether AUC declined. It stops when further removal hurts performance."
)
decision_box(
    "Apply RFECV separately to V3 and Monthly-Only candidate sets",
    "RFECV on V3 (all sources) selected 39 features from 68 candidates. Running it again "
    "on Monthly-Only (51 candidates, no annual features) selected 18 features — and the "
    "resulting model scored higher. This demonstrates that the annual features were not just "
    "low-ranked but actively adding noise: removing them forced the algorithm to find "
    "better monthly substitutes.",
    "The monthly-only model with 18 features outperforms the 39-feature all-sources model. "
    "Fewer features also means faster scoring, less risk of overfitting, and no annual "
    "data lag in production."
)

h2("Model Performance — Final (Monthly-Only, Walk-Forward)")
table(
    ["Task", "Metric", "Score", "Interpretation"],
    [
        ["Binary — predict upward transition",   "AUC (Fold C, 2023)",  "0.865", "Strong discrimination. Random classifier = 0.5; perfect = 1.0."],
        ["3-class — predict direction",          "Macro F1 (Fold C)",   "0.384", "Moderate. The stable class is hardest to distinguish from the others."],
        ["3-class — predict direction",          "Macro F1 (Fold A)",   "0.551", "Higher in 2021 COVID period when up/down signals were more extreme."],
    ]
)
body(
    "The lower Macro F1 on the 3-class task (compared to binary AUC) is expected. The "
    "'stable' class is genuinely ambiguous — a ZIP with moderate signals could reasonably "
    "be classified as slightly up, stable, or slightly down. The binary model's AUC of 0.865 "
    "is the primary performance metric."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9. FORWARD PREDICTIONS & RETROSPECTIVE VALIDATION
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. Forward Predictions & Retrospective Validation")
body(
    "As of March 2026 the project has gone through a full cycle: predictions were made for 2024, "
    "the forward window (Jan–Dec 2025) elapsed, and actual 2025 outcomes became observable. "
    "The model was then retrained on the extended dataset and re-scored on November 2025 features "
    "to generate a new forward signal for Dec 2025–Nov 2026."
)

h2("Retrospective Validation of 2024 Predictions")
body(
    "The original model (trained Jun 2019–Dec 2022) generated predictions for all 2024 months. "
    "With Zillow data through January 2026 and Redfin data through November 2025, it was possible "
    "to compute the actual composite transition score for each ZIP's 12-month forward window and "
    "compare it against the model's predicted probability."
)
table(
    ["Metric", "Value", "Notes"],
    [
        ["AUC — binary (up vs not-up)",    "0.7747", "Solid out-of-distribution generalization"],
        ["Accuracy",                        "0.778",  "75% class imbalance baseline = 0.75"],
        ["F1 — binary (up class)",          "0.489",  "Precision 0.57, Recall 0.42"],
        ["Multi-class Macro F1",            "0.409",  "Down class has very low recall (5%) — model conservative about calling cooling"],
        ["Best cluster — moderate_supply_growing", "AUC 0.822", "Supply-side signals most predictable"],
        ["Weakest cluster — competitive_mid_market", "AUC 0.717", "Mid-market harder to call"],
    ]
)
body(
    "The ~9-point AUC drop from the 2023 held-out test (0.865) is expected: 2024–2025 "
    "represents a post-rate-hike stabilization regime the model was never trained on. "
    "AUC 0.7747 still represents meaningful predictive signal well above random (0.5)."
)

h2("Retrained Model — 2025 Forward Signal")
body(
    "The model was retrained on Jun 2019–Dec 2024 (43,617 labeled rows) incorporating the "
    "validated 2024 transition labels. It was then scored against November 2025 Zillow and "
    "Redfin features to produce the current forward signal."
)
table(
    ["Item", "Value"],
    [
        ["Training range",       "Jun 2019 – Dec 2024"],
        ["Fold D AUC (test 2024)", "0.7240"],
        ["Feature point",        "November 2025"],
        ["Forward window",       "December 2025 – November 2026"],
        ["ZIPs scored",          "650 (5 have Nov 2025 data gaps)"],
        ["Avg opportunity score", "0.183 (lower than 2024 — model calibrated to post-hike market)"],
        ["Verifiable",           "Not until late 2026"],
    ]
)
decision_box(
    "Retrain on extended data rather than keeping the original model",
    "The original model was trained only through 2022. By 2025 it had never seen the 2023 "
    "post-correction stabilization or 2024 dynamics. Retraining on Jun 2019–Dec 2024 gives "
    "the model exposure to two additional years of market behavior including the rate-hike "
    "aftermath, improving its calibration for current conditions.",
    "The retrained model powers the 2025 Signal page in the app. The original Fold C model "
    "is still available in the Validated Map page (2023 held-out results)."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10. THE STREAMLIT APP
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. The Streamlit App")
body(
    "The app has two modes: Investor (default, 5 pages) and Educator (adds 3 methodology "
    "pages). Toggle between them with the mode switcher in the sidebar."
)

h2("Investor Pages")
table(
    ["Page", "What it shows", "When to use it"],
    [
        ["Home",        "Project overview, methodology summary, navigation guide", "Start here — understand the framing before diving into maps"],
        ["2025 Signal", "Forward-looking choropleth map for Dec 2025–Nov 2026 — retrained model on Nov 2025 features.", "Prospecting — identify ZIPs to research further"],
        ["Archetypes",  "Map colored by cluster archetype. Profile cards with opportunity score progress bar.", "Understand the structural character of a market before reading its prediction"],
        ["Compare",     "Side-by-side 2–4 ZIP comparison: opportunity score, net direction, SHAP drivers, summary table", "Narrow down from a shortlist — compare specific candidates head-to-head"],
        ["Deep Dive",   "Single ZIP: monthly opportunity score over time, key feature trends vs. cluster average, SHAP waterfall by month", "Deep analysis of one ZIP — understand why the model scored it and how it evolved"],
    ]
)

h2("Educator Pages (toggle Educator mode in sidebar)")
table(
    ["Page", "What it shows"],
    [
        ["Validated Map",     "2023 test-set choropleth — same map as 2024 Signal but for validated data. Shows opportunity score or net direction."],
        ["Model Comparison",  "V3 (all sources) vs. Monthly-Only head-to-head: metrics, fold breakdown, side-by-side maps, feature lists."],
        ["Performance",       "Walk-forward fold metrics, version history table (V1→V2→V3→Monthly-Only), top-15 features by mean |SHAP| importance."],
    ]
)

h2("Sidebar Controls")
table(
    ["Control", "Effect"],
    [
        ["Investor / Educator mode", "Shows or hides the 3 methodology pages in the top navigation"],
        ["State filter",             "Filter all maps and tables to NC only, SC only, or both"],
        ["Cluster archetype filter", "Limit the map and tables to specific cluster types"],
        ["Opportunity threshold",    "Highlight bar — ZIPs above this % are called 'high opportunity' in tables and metrics"],
    ]
)

h2("Reading SHAP Waterfall Charts (Deep Dive page)")
body(
    "SHAP (SHapley Additive exPlanations) values explain why the model gave a specific ZIP "
    "a specific score in a specific month. Each bar represents one feature's contribution "
    "to pushing the prediction higher (red, positive) or lower (blue, negative)."
)
table(
    ["SHAP concept", "Plain English"],
    [
        ["Base value",         "The model's average prediction across all ZIPs — the starting point before any ZIP-specific features are considered"],
        ["Positive SHAP (red)", "This feature's value for this ZIP, in this month, pushed the prediction toward upward transition"],
        ["Negative SHAP (blue)", "This feature's value pushed the prediction toward stable or cooling"],
        ["Magnitude",          "The longer the bar, the more that feature influenced this specific prediction"],
        ["Sum",                "All SHAP values plus the base value equals the model's final raw score (log-odds) for this ZIP/month"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11. KNOWN LIMITATIONS
# ═══════════════════════════════════════════════════════════════════════════════
h1("11. Known Limitations")
body("These are important to acknowledge honestly, especially in a presentation context.")

table(
    ["Limitation", "Why it exists", "How to address it in conversation"],
    [
        ["ZIP code aggregation",
         "The model operates at ZIP level — a ZIP can contain very different block-level markets",
         "Position as a screening tool. Deep Dive identifies ZIPs worth researching further at a finer level."],
        ["Engineered targets (no ground truth)",
         "There is no official dataset of 'neighborhood transitions'. Our target is a proxy composite score.",
         "The composite is grounded in three real market signals. The model predicts the composite accurately, not an arbitrary label."],
        ["Clustering silhouette = 0.16",
         "Real estate markets are a continuum — ZIPs don't fall into discrete buckets",
         "Present clusters as 'market tendencies on a spectrum', not hard categories. Using the same 18 model features improved silhouette from 0.11 to 0.16; profiles remain interpretable even with fuzzy boundaries."],
        ["2022 AUC dip (0.815)",
         "Fed rate-hike cycle created conditions outside the training distribution",
         "Acknowledge it directly. The recovery to 0.865 in 2023 is the more important signal — shows durability."],
        ["2025 predictions are unverified",
         "The 12-month outcome window for 2025 predictions ends in late 2026",
         "Be explicit. The 2024 predictions were validated at AUC 0.7747 on actual 2025 outcomes — that gives confidence, but 2025 predictions still cannot be verified until late 2026."],
        ["Annual data lag",
         "ACS/IRS data is published 12–24 months after the reference year",
         "This was the primary reason we dropped annual features from the winning model. The monthly-only model has no data-lag risk."],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12. FREQUENTLY ASKED QUESTIONS
# ═══════════════════════════════════════════════════════════════════════════════
h1("12. Anticipated Questions & Answers")

h2("Why not use a neural network / deep learning model?")
body(
    "Gradient Boosting on tabular data consistently outperforms deep learning when the dataset "
    "is relatively small (20,000 rows), features are structured numerical inputs, and "
    "interpretability matters. SHAP values give us feature-level explanations that would be "
    "much harder to produce from a neural network. We also only have ~2.5 years of labeled "
    "data — neural networks need more to generalize well."
)

h2("Why not use LSTM or time-series-specific models?")
body(
    "LSTMs and similar sequence models require long, clean time series per entity. Our panel "
    "starts in mid-2019 and ends in late 2021 for labeled data — roughly 30 months per ZIP. "
    "That is too short for an LSTM to learn meaningful temporal dependencies beyond what "
    "lag features already capture. Our engineered lag/rolling features provide the model with "
    "the same temporal context more reliably and interpretably."
)

h2("Why 12-month prediction horizon?")
body(
    "A 12-month horizon is the natural decision cycle for real estate investors — it aligns "
    "with typical due diligence and acquisition timelines. Shorter horizons (3–6 months) "
    "would be noisier and harder to act on. Longer horizons (24+ months) would require more "
    "data than we have for the forward window."
)

h2("How do you know the model isn't overfitting?")
body(
    "Walk-forward validation prevents overfitting by construction: the model never sees "
    "future data during training. We also use RFECV to select features — this removes "
    "low-signal features that would otherwise allow the model to fit noise. The consistency "
    "of AUC across folds A (0.844) and C (0.865) — despite very different market conditions — "
    "is evidence of genuine generalization, not overfitting."
)

h2("What does a 'high opportunity score' mean in practice?")
body(
    "A score of 0.70 means the model assigns a 70% probability that this ZIP will enter an "
    "upward housing transition over the next 12 months — defined as its composite score of "
    "home value growth, inventory tightening, and market competitiveness exceeding the 75th "
    "percentile of all ZIPs in that same time window. It is a relative ranking signal, not "
    "an absolute forecast of price appreciation percentage."
)

h2("Is this project saying you SHOULD invest in high-opportunity ZIPs?")
body(
    "No. The tool is an early-warning and screening system, not investment advice. A high "
    "opportunity score identifies ZIPs that the model predicts are likely to enter transition "
    "conditions — it does not account for property-specific factors, local regulation, "
    "financing conditions, or the investor's specific strategy. It narrows the search space; "
    "it does not replace due diligence."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13. TECHNICAL SETUP
# ═══════════════════════════════════════════════════════════════════════════════
h1("13. Technical Setup & Reproducibility")

h2("Notebook Pipeline — Run in This Order")
table(
    ["Step", "Notebook / Script", "Output", "Notes"],
    [
        ["1", "exploration.ipynb",           "consolidated_zip_month_redfin_gap_filtered_model_cols.csv", "Data quality audit, ZIP coverage filter"],
        ["2", "data_prep_and_eda.ipynb",     "modeling_ready.csv",                                        "Consolidation and EDA"],
        ["3", "feature_engineering.ipynb",   "features_engineered.csv",                                   "Lag, rolling, acceleration, baseline features"],
        ["4", "target_engineering.ipynb",    "modeling_dataset.csv",                                      "Composite score + binary/3-class targets"],
        ["5", "clustering.ipynb",            "modeling_dataset_with_clusters.csv",                        "K-means k=5 archetypes"],
        ["6", "modeling.ipynb",              "model_predictions_v3.csv, shap_values.csv",                 "V3 model (all sources)"],
        ["7", "scripts/run_monthly_model.py","model_predictions_monthly.csv, shap_values_monthly.csv, model_comparison.json", "Monthly-only model (winner)"],
        ["8", "scripts/generate_forward_predictions.py", "model_predictions_forward_2024.csv, model_predictions_multiclass_forward_2024.csv", "2024 forward signal"],
    ]
)

h2("Package Management")
body(
    "The project uses uv for dependency management with pyproject.toml + uv.lock for "
    "reproducibility. Always use 'uv add <package>' to install dependencies — never pip install "
    "directly, as this bypasses the lock file."
)
body(
    "Exception: shap must be installed with 'uv pip install shap' due to a cross-platform "
    "numba version conflict that prevents it from being added to the lock file. This is only "
    "needed for running the modeling notebooks — the Streamlit app does not import shap "
    "(SHAP values are pre-computed and stored as CSVs)."
)

h2("Deployment")
body(
    "The app is deployed on Hugging Face Spaces (Streamlit SDK) at "
    "https://huggingface.co/spaces/timagonch/nc-sc-market-monitor. "
    "Only the 5 runtime data files (~57MB total) are included in the Space repository. "
    "The large intermediate pipeline files are gitignored and remain local."
)

# ── Save ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
